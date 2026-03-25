"""
app.py - Flask backend for Automated Requirements Analysis Using NLP
"""

import io
import textwrap
from datetime import datetime
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from analyzer import analyze_document
from requirement_cleaner import clean_requirements
from requirement_refiner import refine_requirements
from srs_generator import generate_srs_content, generate_fallback_content

app = Flask(__name__)
CORS(app, resources={r"/api/*": {"origins": "*"}})

ALLOWED_EXTENSIONS = {"txt", "docx"}
MAX_FILE_SIZE_MB = 5

_TOOL_ONLY_TERMS = [
    "nlp", "natural language processing", "requirement extraction", "srs generation",
    "quality scoring", "quality assessment", "analysis dashboard", "analyzes documents",
    "analyze documents",
]


def _is_tool_only_content(text: str) -> bool:
    lo = (text or "").lower()
    return any(term in lo for term in _TOOL_ONLY_TERMS)


def _sanitize_ai_content(ai_content: dict) -> dict:
    """Remove analysis-tool capabilities from generated section prose."""
    cleaned = {}
    for key, value in (ai_content or {}).items():
        if isinstance(value, str):
            kept_lines = [ln for ln in value.splitlines() if not _is_tool_only_content(ln)]
            cleaned[key] = "\n".join(kept_lines).strip()
        else:
            cleaned[key] = value
    return cleaned


def _filter_target_system_requirements(refined: list) -> list:
    """
    Keep only requirements that belong to the target system boundary.
    """
    filtered = []
    for req in refined or []:
        text = req.get("text", "") or req.get("sentence", "")
        if _is_tool_only_content(text):
            continue
        filtered.append(req)
    return filtered


def _req_text(req: dict) -> str:
    """Return requirement text across raw/refined payload shapes."""
    return req.get("sentence") or req.get("text") or req.get("description") or ""


def allowed_file(filename: str) -> bool:
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract plain text from a .docx file using python-docx."""
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n".join(paragraphs)
    except Exception as e:
        raise ValueError(f"Could not parse .docx file: {str(e)}")


@app.errorhandler(413)
def request_entity_too_large(error):
    return jsonify({"error": f"File too large. Maximum allowed size is {MAX_FILE_SIZE_MB}MB."}), 413


@app.route("/api/health", methods=["GET"])
def health():
    return jsonify({"status": "ok", "message": "ARAQAT backend is running."})


@app.route("/api/analyze", methods=["POST"])
def analyze():
    if "file" not in request.files:
        return jsonify({"error": "No file part in the request."}), 400

    file = request.files["file"]

    if file.filename == "":
        return jsonify({"error": "No file selected."}), 400

    if not allowed_file(file.filename):
        ext = file.filename.rsplit(".", 1)[-1].upper() if "." in file.filename else "unknown"
        return jsonify({
            "error": f"Unsupported file format (.{ext.lower()}). Please upload a .txt or .docx file."
        }), 400

    file_bytes = file.read()
    ext = file.filename.rsplit(".", 1)[1].lower()

    try:
        if ext == "docx":
            raw_text = extract_text_from_docx(file_bytes)
        else:
            raw_text = file_bytes.decode("utf-8", errors="ignore")
    except ValueError as e:
        return jsonify({"error": str(e)}), 400
    except Exception as e:
        return jsonify({"error": f"Could not read file: {str(e)}"}), 500

    if not raw_text.strip():
        return jsonify({"error": "The uploaded file appears to be empty."}), 400

    try:
        result = analyze_document(raw_text)
    except Exception as e:
        return jsonify({"error": f"Analysis failed: {str(e)}"}), 500

    return jsonify(result), 200


# ─────────────────────────────────────────────────────────────────────────────
# IEEE 830-1998 SRS Export  (Karl E. Wiegers template)
# ─────────────────────────────────────────────────────────────────────────────

def _categorise_nfr(sentence: str) -> str:
    """Map an NFR sentence to one of the IEEE 830 §5 sub-categories."""
    lo = sentence.lower()
    PERF  = {"performance","response time","latency","throughput","uptime","load","concurrent","benchmark"}
    SAFE  = {"safety","harm","damage","loss","safeguard","hazard","fault","failure","prevention"}
    SEC   = {"security","encrypt","auth","privacy","compliance","audit","certificate","access control","vulnerability"}
    QUAL  = {"reliability","maintainability","availability","usability","portability","testability",
             "reusability","interoperability","scalability","flexibility","robustness","adaptability"}
    if any(k in lo for k in SAFE): return "safety"
    if any(k in lo for k in SEC):  return "security"
    if any(k in lo for k in PERF): return "performance"
    if any(k in lo for k in QUAL): return "quality"
    return "other"


# ─────────────────────────────────────────────────────────────────────────────
# Smart SRS content helpers — generate meaningful text from requirements
# ─────────────────────────────────────────────────────────────────────────────

def _req_text(req: dict) -> str:
    """
    Extract the requirement text string from a requirement dict.
    Handles key variants from different pipeline stages:
      - 'text'        (from requirement_refiner)
      - 'sentence'    (from analyzer)
      - 'description' (from older fallback dicts)
    """
    return (
        req.get("text") or
        req.get("sentence") or
        req.get("description") or
        ""
    )


def _first_n_words(sentence: str, n: int = 12) -> str:
    words = sentence.split()
    snippet = " ".join(words[:n])
    return (snippet + "…") if len(words) > n else snippet


def _infer_scope(project: str, frs: list, nfrs: list) -> str:
    """
    Auto-generate a Product Scope paragraph by scanning requirement keywords.
    """
    all_text = " ".join(_req_text(r).lower() for r in frs + nfrs)

    # Theme detection
    themes = []
    THEME_MAP = [
        (["user", "login", "register", "account", "profile", "authentication"], "user account management"),
        (["report", "analytics", "statistics", "chart", "graph"],                 "data reporting and analytics"),
        (["payment", "invoice", "billing", "transaction", "checkout", "order"],   "payment and transaction processing"),
        (["search", "filter", "query", "browse", "find"],                          "search and discovery"),
        (["notification", "alert", "email", "sms", "message", "send"],             "user notifications"),
        (["upload", "download", "file", "document", "export", "import"],           "file and document management"),
        (["admin", "manage", "configure", "setting", "role", "permission"],        "administration and configuration"),
        (["data", "store", "database", "record", "save",  "retrieve"],             "data storage and retrieval"),
        (["api", "integration", "interface", "service", "connect", "third-party"], "external system integration"),
        (["security", "encrypt", "access control", "privacy"],                     "security and access control"),
    ]
    for keywords, label in THEME_MAP:
        if any(k in all_text for k in keywords):
            themes.append(label)

    fr_cnt  = len(frs)
    nfr_cnt = len(nfrs)

    if themes:
        theme_str = ", ".join(themes[:4])
        if len(themes) > 4:
            theme_str += f", and {len(themes) - 4} additional capability area(s)"
        scope = (
            f"{project} is a software system designed to support {theme_str}. "
            f"The system delivers {fr_cnt} functional capabilities and must satisfy "
            f"{nfr_cnt} non-functional quality attributes as defined in this specification."
        )
    else:
        scope = (
            f"{project} is a software system comprising {fr_cnt} functional requirements "
            f"and {nfr_cnt} non-functional requirements. "
            f"The system provides core operational capabilities for its intended users."
        )
    return scope


def _infer_functions(frs: list, max_items: int = 8) -> list:
    """
    Return a bullet list of key capability summaries derived from FR sentences.
    """
    if not frs:
        return ["(No functional requirements identified.)"]
    bullets = []
    for i, req in enumerate(frs[:max_items]):
        bullets.append(f"• FR-{i+1:03d}: {_first_n_words(_req_text(req), 14)}")
    if len(frs) > max_items:
        bullets.append(f"• … and {len(frs) - max_items} additional functional requirements (see §4 for full list).")
    return bullets


def _infer_user_classes(frs: list, nfrs: list) -> list:
    """
    Detect user roles from requirement text.
    """
    all_text = " ".join(_req_text(r).lower() for r in frs + nfrs)
    classes = [
        ("End Users",
         "Primary users who interact with the system's core features and functions. "
         "Expected to have standard technology literacy."),
    ]
    if any(k in all_text for k in ["admin", "administrator", "manage", "role", "permission", "configure"]):
        classes.append((
            "Administrators",
            "System administrators responsible for managing user accounts, system "
            "configuration, access control, and operational settings."
        ))
    if any(k in all_text for k in ["manager", "supervisor", "approver", "reviewer"]):
        classes.append((
            "Managers / Reviewers",
            "Users with elevated privileges who review, approve, or oversee activities "
            "performed by other users."
        ))
    if any(k in all_text for k in ["guest", "anonymous", "public", "visitor", "unauthenticated"]):
        classes.append((
            "Guest / Public Users",
            "Unauthenticated users with limited read-only or exploratory access to the system."
        ))
    if any(k in all_text for k in ["api", "integration", "external system", "third-party", "webhook"]):
        classes.append((
            "External Systems / API Consumers",
            "Third-party applications or services that interact programmatically via defined APIs."
        ))
    return classes


def _infer_environment(frs: list, nfrs: list) -> list:
    """
    Generate an operating environment description from requirement keywords.
    """
    all_text = " ".join(_req_text(r).lower() for r in frs + nfrs)
    env = []

    # Platform
    if any(k in all_text for k in ["mobile", "ios", "android", "app", "smartphone"]):
        env.append("Mobile Application: iOS (14+) and Android (10+) via native or hybrid app.")
    if any(k in all_text for k in ["web", "browser", "html", "http", "url", "page", "portal", "interface"]):
        env.append("Web Browser: Compatible with Chrome 110+, Firefox 110+, Safari 16+, and Edge 110+.")
    if not env:  # default to web
        env.append("Web Browser: Compatible with Chrome 110+, Firefox 110+, Safari 16+, and Edge 110+.")

    # Backend
    env.append("Server-Side: Backend application server (RESTful API) hosted in a cloud or on-premise environment.")

    # Database
    if any(k in all_text for k in ["database", "db", "store", "record", "sql", "data"]):
        env.append("Database: Relational or document-oriented database for persistent data storage.")

    # Network
    env.append("Network: Requires stable internet connectivity (minimum 5 Mbps for optimal performance).")

    # Security
    if any(k in all_text for k in ["encrypt", "https", "tls", "ssl", "secure"]):
        env.append("Security: HTTPS with TLS 1.2+ for all client-server communications.")

    return env


def _infer_interfaces(frs: list, nfrs: list, project: str) -> dict:
    """
    Generate §3 External Interface descriptions based on requirement keywords.
    """
    all_text = " ".join(_req_text(r).lower() for r in frs + nfrs)
    ifaces = {}

    # User Interfaces
    ui_desc = (
        f"The {project} system shall provide a responsive graphical user interface "
        f"accessible via web browser. Interfaces include navigation, data entry forms, "
        f"and feedback dialogs. Designs shall follow WCAG 2.1 AA "
        f"accessibility guidelines."
    )
    ifaces["ui"] = ui_desc

    # Hardware
    hw_desc = (
        "Client hardware requires a device capable of running a modern web browser with "
        "minimum 2 GB RAM. Server hardware shall support the expected concurrent user load "
        "as defined in §5.1."
    )
    ifaces["hw"] = hw_desc

    # Software
    sw_parts = [f"The {project} system interacts with the following software components:"]
    if any(k in all_text for k in ["database", "db", "sql", "store", "record"]):
        sw_parts.append("• Database engine (e.g., PostgreSQL, MySQL, or MongoDB) for persistent storage.")
    if any(k in all_text for k in ["email", "smtp", "notification", "alert"]):
        sw_parts.append("• Email service provider (e.g., SendGrid, AWS SES) for outbound notifications.")
    if any(k in all_text for k in ["api", "rest", "integration", "third-party", "external"]):
        sw_parts.append("• External REST APIs for third-party data or service integration.")
    if any(k in all_text for k in ["payment", "billing", "transaction"]):
        sw_parts.append("• Payment gateway (e.g., Stripe, PayPal) for financial transactions.")
    if len(sw_parts) == 1:
        sw_parts.append("• Standard OS libraries and runtime environment.")
    ifaces["sw"] = "\n    ".join(sw_parts)

    # Communications
    comm_parts = ["Communication protocols and standards:"]
    comm_parts.append("• HTTPS / TLS 1.2+ for all client-server data transfer.")
    if any(k in all_text for k in ["api", "rest", "json", "xml", "webhook"]):
        comm_parts.append("• RESTful JSON APIs for service-to-service communication.")
    if any(k in all_text for k in ["real-time", "websocket", "live", "push", "streaming"]):
        comm_parts.append("• WebSocket connections for real-time data streaming.")
    if any(k in all_text for k in ["email", "smtp", "notification"]):
        comm_parts.append("• SMTP for transactional email delivery.")
    ifaces["comm"] = "\n    ".join(comm_parts)

    return ifaces



def _build_ieee_srs(requirements, metrics, meta, ai=None) -> str:
    """
    Build a full IEEE 830-1998 SRS text document.
    meta keys: project_name, author, organization, version, date_created
    ai: dict from srs_generator with rich section prose (or None for heuristic fallback)
    """
    if ai is None:
        ai = {}

    project  = meta.get("project_name", "<Project>")
    author   = meta.get("author",       "<author>")
    org      = meta.get("organization", "<organization>")
    version  = meta.get("version",      "1.0")
    date     = meta.get("date_created") or datetime.now().strftime("%Y-%m-%d")

    frs  = [r for r in requirements if r.get("type") == "FR"]
    nfrs = [r for r in requirements if r.get("type") == "NFR"]

    perf_nfrs   = [r for r in nfrs if _categorise_nfr(_req_text(r)) == "performance"]
    safety_nfrs = [r for r in nfrs if _categorise_nfr(_req_text(r)) == "safety"]
    sec_nfrs    = [r for r in nfrs if _categorise_nfr(_req_text(r)) == "security"]
    qual_nfrs   = [r for r in nfrs if _categorise_nfr(_req_text(r)) == "quality"]
    other_nfrs  = [r for r in nfrs if _categorise_nfr(_req_text(r)) == "other"]

    total   = metrics.get("total_requirements", 0)
    fr_cnt  = metrics.get("fr_count",  0)
    nfr_cnt = metrics.get("nfr_count", 0)
    v_cnt   = metrics.get("vague_count", 0)
    score   = metrics.get("quality_score", 0)
    grade   = ("A (Excellent)" if score >= 90 else
               "B (Good)"      if score >= 75 else
               "C (Fair)"      if score >= 55 else "D (Needs Work)")

    vague_reqs = [r for r in requirements if r.get("is_vague")]

    W = 70  # page width
    SEP  = "=" * W
    SEP2 = "-" * W

    def centre(text): return text.center(W)
    def h1(num, title): return f"\n\n{num}  {title}\n{SEP2}"
    def h2(num, title): return f"\n{num}  {title}"
    def h3(num, title): return f"\n    {num}  {title}"
    def p(text):  return f"\n    {text}" if text else ""
    def flag(req): return f"  *** VAGUE: {', '.join(req['vague_words'])} ***" if req.get("vague_words") else ""

    L = []

    # ── COVER PAGE ───────────────────────────────────────────────────────────
    L += [
        "", SEP,
        centre("Software Requirements Specification"),
        "", centre("for"), "",
        centre(project), "",
        centre(f"Version {version} approved"), "",
        centre(f"Prepared by {author}"), "",
        centre(org), "",
        centre(date), "",
        SEP, "",
        " " * 4 + "Copyright © 1999 by Karl E. Wiegers. "
                  "Permission is granted to use, modify, and distribute this document.",
        SEP,
    ]

    # ── TABLE OF CONTENTS ────────────────────────────────────────────────────
    TOC = [
        ("Table of Contents",                          "i"),
        ("Revision History",                           "ii"),
        ("1.  Introduction",                           "1"),
        ("    1.1  Purpose",                           "1"),
        ("    1.2  Document Conventions",              "1"),
        ("    1.3  Intended Audience and Reading Suggestions", "1"),
        ("    1.4  Product Scope",                     "1"),
        ("    1.5  References",                        "1"),
        ("2.  Overall Description",                    "2"),
        ("    2.1  Product Perspective",               "2"),
        ("    2.2  Product Functions",                 "2"),
        ("    2.3  User Classes and Characteristics",  "2"),
        ("    2.4  Operating Environment",             "2"),
        ("    2.5  Design and Implementation Constraints", "2"),
        ("    2.6  User Documentation",                "2"),
        ("    2.7  Assumptions and Dependencies",      "3"),
        ("3.  External Interface Requirements",        "3"),
        ("    3.1  User Interfaces",                   "3"),
        ("    3.2  Hardware Interfaces",               "3"),
        ("    3.3  Software Interfaces",               "3"),
        ("    3.4  Communications Interfaces",         "3"),
        ("4.  System Features (Functional Requirements)", "4"),
        ("5.  Other Nonfunctional Requirements",       "5"),
        ("    5.1  Performance Requirements",          "5"),
        ("    5.2  Safety Requirements",               "5"),
        ("    5.3  Security Requirements",             "5"),
        ("    5.4  Software Quality Attributes",       "5"),
        ("    5.5  Business Rules / Other NFRs",       "5"),
        ("6.  Other Requirements",                     "5"),
        ("Appendix A: Glossary",                       "6"),
        ("Appendix B: System Models",                  "6"),
        ("Appendix C: To Be Determined List",          "6"),
    ]
    L.append("\n\nTable of Contents\n" + SEP2)
    for entry, pg in TOC:
        dots = "." * max(2, W - len(entry) - len(pg) - 2)
        L.append(f"  {entry}{dots}{pg}")

    # ── REVISION HISTORY ─────────────────────────────────────────────────────
    L.append(f"\n\nRevision History\n{SEP2}")
    L.append(f"  {'Name':<20} {'Date':<14} {'Reason For Changes':<26} Version")
    L.append(f"  {'─'*20} {'─'*14} {'─'*26} {'─'*7}")
    L.append(f"  {'Generated':<20} {date:<14} {'Initial SRS':<26} {version}")

    # ── §1  INTRODUCTION ─────────────────────────────────────────────────────
    L.append(h1("1.", "Introduction"))
    L.append(h2("1.1", "Purpose"))
    L.append(p(f"This Software Requirements Specification (SRS) describes the requirements"))
    L.append(p(f"for {project}. It was generated from the provided requirements document "))
    L.append(p(f"(version {version}, {date})."))

    L.append(h2("1.2", "Document Conventions"))
    L.append(p("FR-XXX  = Functional Requirement  |  NFR-XXX = Non-Functional Requirement"))
    L.append(p("Statements marked *** VAGUE *** contain ambiguous language and must be refined."))
    L.append(p("All requirements carry equal priority unless explicitly annotated."))

    L.append(h2("1.3", "Intended Audience and Reading Suggestions"))
    L.append(p("Intended for developers, project managers, testers, and QA staff."))
    L.append(p("Recommendation: read §2 for context, then §4 and §5 for specific requirements."))

    L.append(h2("1.4", "Product Scope"))
    scope_text = ai.get("product_scope") or _infer_scope(project, frs, nfrs)
    for line in textwrap.wrap(scope_text, 65):
        L.append(p(line))

    L.append(h2("1.5", "References"))
    L.append(p("• IEEE Std 830-1998, IEEE Recommended Practice for Software Requirements Specifications."))
    L.append(p("• Karl E. Wiegers, Software Requirements, 2nd Edition. Microsoft Press."))
    L.append(p(f"• Requirements Specification Source Document, {date}."))

    # ── §2  OVERALL DESCRIPTION ──────────────────────────────────────────────
    L.append(h1("2.", "Overall Description"))
    L.append(h2("2.1", "Product Perspective"))
    perspective = ai.get("product_perspective") or scope_text
    for line in textwrap.wrap(perspective, 65):
        L.append(p(line))
    L.append(p(f"This SRS documents {total} requirements ({fr_cnt} FR, {nfr_cnt} NFR) "
               f"covering all aspects of the {project} system."))

    L.append(h2("2.2", "Product Functions"))
    if ai.get("product_functions"):
        for line in textwrap.wrap(ai["product_functions"], 65):
            L.append(p(line))
    else:
        L.append(p(f"Key capabilities of {project} (see §4 for full detail):"))
        for bullet in _infer_functions(frs):
            L.append(p(f"  {bullet}"))

    L.append(h2("2.3", "User Classes and Characteristics"))
    if ai.get("user_classes"):
        for line in textwrap.wrap(ai["user_classes"], 65):
            L.append(p(line))
    else:
        for role, desc in _infer_user_classes(frs, nfrs):
            L.append(p(f"  {role}:"))
            for line in textwrap.wrap(desc, 60):
                L.append(p(f"    {line}"))

    L.append(h2("2.4", "Operating Environment"))
    if ai.get("operating_environment"):
        for line in textwrap.wrap(ai["operating_environment"], 65):
            L.append(p(line))
    else:
        for item in _infer_environment(frs, nfrs):
            L.append(p(f"  • {item}"))

    L.append(h2("2.5", "Design and Implementation Constraints"))
    if ai.get("design_constraints"):
        for line in textwrap.wrap(ai["design_constraints"], 65):
            L.append(p(line))
    else:
        L.append(p("• The system must comply with IEEE 830-1998 specification structure."))
        L.append(p("• All inter-system communication must use HTTPS with TLS 1.2 or higher."))
        L.append(p(f"• {v_cnt} vague requirement(s) require clarification before implementation. See Appendix C."))
        if v_cnt > 0:
            L.append(p("• Ambiguous requirements should be resolved before development milestones."))

    L.append(h2("2.6", "User Documentation"))
    if ai.get("user_documentation"):
        for line in textwrap.wrap(ai["user_documentation"], 65):
            L.append(p(line))
    else:
        L.append(p(f"• {project} User Manual: step-by-step instructions for end users."))
        L.append(p("• System Administrator Guide: installation, configuration, and maintenance."))
        L.append(p("• API Reference: endpoint documentation for integration developers."))
        L.append(p("• Release Notes: per-version change log and known issues."))

    L.append(h2("2.7", "Assumptions and Dependencies"))
    if ai.get("assumptions_dependencies"):
        for line in textwrap.wrap(ai["assumptions_dependencies"], 65):
            L.append(p(line))
    else:
        L.append(p("• The target deployment environment meets the specifications in §2.4."))
        if any(k in " ".join(_req_text(r).lower() for r in frs+nfrs)
               for k in ["third-party", "external", "api", "service"]):
            L.append(p("• Third-party services and APIs maintain their documented SLA."))
        L.append(p("• Requirements marked VAGUE will be clarified before sprint planning."))

    # ── §3  EXTERNAL INTERFACE REQUIREMENTS ──────────────────────────────────
    ifaces = _infer_interfaces(frs, nfrs, project)
    L.append(h1("3.", "External Interface Requirements"))

    L.append(h2("3.1", "User Interfaces"))
    ui_text = ai.get("user_interfaces") or ifaces["ui"]
    for line in textwrap.wrap(ui_text, 65):
        L.append(p(line))

    L.append(h2("3.2", "Hardware Interfaces"))
    hw_text = ai.get("hardware_interfaces") or ifaces["hw"]
    for line in textwrap.wrap(hw_text, 65):
        L.append(p(line))

    L.append(h2("3.3", "Software Interfaces"))
    sw_text = ai.get("software_interfaces") or ifaces["sw"]
    for line in sw_text.split("\n"):
        L.append(p(line))

    L.append(h2("3.4", "Communications Interfaces"))
    comm_text = ai.get("communications_interfaces") or ifaces["comm"]
    for line in comm_text.split("\n"):
        L.append(p(line))

    # ── §4  SYSTEM FEATURES (Functional Requirements) ────────────────────────
    L.append(h1("4.", "System Features (Functional Requirements)"))
    L.append(p(f"This section lists all {fr_cnt} functional requirements identified in the document."))
    L.append(p("Each feature corresponds to a requirement from the source document."))

    refined_frs = ai.get("refined_functional_requirements", [])
    if not refined_frs:
        refined_frs = [{"original_id": r.get("id"), "description": _req_text(r), "stimulus_response": "<Describe user actions...>"} for r in frs]

    if refined_frs:
        for i, req in enumerate(refined_frs, 1):
            fid = f"4.{i}"
            L.append(f"\n{h2(fid, f'FR-{i:03d}')}")
            L.append(h3(f"{fid}.1", "Description and Priority"))
            L.append(p(f"        {req.get('description', '')}"))
            L.append(p(f"        Priority: HIGH"))
            L.append(h3(f"{fid}.2", "Stimulus/Response Sequences"))
            L.append(p(f"        {req.get('stimulus_response', '')}"))
            L.append(h3(f"{fid}.3", "Functional Requirements"))
            L.append(p(f"        REQ-{i:03d}: {req.get('description', '')}"))
    else:
        L.append(p("(No functional requirements were identified in the source document.)"))

    # ── §5  OTHER NONFUNCTIONAL REQUIREMENTS ─────────────────────────────────
    L.append(h1("5.", "Other Nonfunctional Requirements"))

    def nfr_block(sec_num, title, subset, tag):
        L.append(h2(sec_num, title))
        if subset:
            for j, r in enumerate(subset, 1):
                desc = r.get("description", r.get("sentence", ""))
                L.append(p(f"    {tag}-{j:03d}: {desc}"))
        else:
            L.append(p("    >>> [PLACEHOLDER] No requirements found for this category. <<<"))
            L.append(p("    >>> Please add relevant requirements or mark as Not Applicable. <<<"))

    refined_nfrs = ai.get("refined_non_functional_requirements", [])
    if not refined_nfrs:
        refined_nfrs = [{"original_id": r.get("id"), "description": _req_text(r), "category": "Other"} for r in nfrs]
        
    perf_nfrs   = [r for r in refined_nfrs if "performance" in r.get("category", "").lower()]
    sec_nfrs    = [r for r in refined_nfrs if "security" in r.get("category", "").lower()]
    usab_nfrs   = [r for r in refined_nfrs if "usability" in r.get("category", "").lower()]
    scal_nfrs   = [r for r in refined_nfrs if "scalability" in r.get("category", "").lower()]
    reli_nfrs   = [r for r in refined_nfrs if "reliability" in r.get("category", "").lower()]

    nfr_block("5.1", "Performance Requirements",  perf_nfrs, "PERF")
    nfr_block("5.2", "Security Requirements",     sec_nfrs,  "SEC")
    nfr_block("5.3", "Usability Requirements",    usab_nfrs, "USAB")
    nfr_block("5.4", "Scalability Requirements",  scal_nfrs, "SCAL")
    nfr_block("5.5", "Reliability Requirements",  reli_nfrs, "REL")

    # ── §6  OTHER REQUIREMENTS ───────────────────────────────────────────────
    L.append(h1("6.", "Other Requirements"))
    if ai.get("other_requirements"):
        for line in textwrap.wrap(ai["other_requirements"], 65):
            L.append(p(line))
    else:
        all_lo = " ".join(_req_text(r).lower() for r in frs + nfrs)
        if any(k in all_lo for k in ["language", "locale", "internation", "multilingual", "i18n", "region"]):
            L.append(p("Internationalisation: The system shall support multiple locales/languages "
                       "as specified during project initiation."))
        if any(k in all_lo for k in ["gdpr", "legal", "compliance", "regulation", "law", "privacy policy"]):
            L.append(p("Legal/Regulatory: The system must comply with applicable data protection "
                       "regulations (e.g., GDPR, HIPAA) as required by jurisdiction."))
        if any(k in all_lo for k in ["database", "backup", "archive", "retention", "purge"]):
            L.append(p("Data Retention: Data shall be retained for a minimum period as defined "
                       "by applicable policy; purged securely thereafter."))
        else:
            L.append(p("Database: Data persistence and storage requirements are detailed in §3.3."))
            L.append(p("Legal: The system shall comply with applicable laws and regulations in the "
                       "target deployment region."))

    # ── APPENDIX A: GLOSSARY ─────────────────────────────────────────────────
    L.append(f"\n\nAppendix A: Glossary\n{SEP2}")
    L.append(p("FR    — Functional Requirement: a capability the system must perform."))
    L.append(p("NFR   — Non-Functional Requirement: a quality/constraint the system must meet."))
    L.append(p("VAGUE — Ambiguous language; statement needs a quantifiable, measurable form."))
    L.append(p("TBD   — To Be Determined; information not yet available."))
    L.append(p("SRS   — Software Requirements Specification."))
    L.append(p("IEEE  — Institute of Electrical and Electronics Engineers."))
    if any(k in " ".join(r.get("sentence", r.get("description", "")).lower() for r in frs+nfrs) for k in ["api", "rest"]):
        L.append(p("API   — Application Programming Interface: a set of defined interaction rules."))
    if any(k in " ".join(r.get("sentence", r.get("description", "")).lower() for r in frs+nfrs) for k in ["authenticate", "auth", "login"]):
        L.append(p("Auth  — Authentication/Authorization: verifying identity and access rights."))

    # ── APPENDIX B: ANALYSIS MODELS ──────────────────────────────────────────
    L.append(f"\n\nAppendix B: System Models\n{SEP2}")
    L.append(p(f"System Models for {project}"))
    L.append(p(""))

    # B.1 System Overview
    L.append(p("  B.1  System Overview"))
    all_text_lo = " ".join(_req_text(r).lower() for r in frs + nfrs)
    themes_b = []
    theme_map_b = [
        (["upload", "import", "file", "document"], "document ingestion and processing"),
        (["login", "auth", "register", "account"], "user authentication and account management"),
        (["search", "filter", "query", "browse"], "search and data retrieval"),
        (["report", "analytics", "chart"], "reporting and analytics"),
        (["export", "generate", "download"], "output and data export"),
        (["notify", "alert", "email", "message"], "notifications and messaging"),
        (["payment", "invoice", "billing"], "payment processing"),
        (["admin", "manage", "configure", "role"], "administration and configuration"),
    ]
    for keywords_b, label_b in theme_map_b:
        if any(k in all_text_lo for k in keywords_b):
            themes_b.append(label_b)
    if themes_b:
        L.append(p(f"  The {project} system provides the following core capabilities:"))
        for tb in themes_b:
            L.append(p(f"    - {tb.title()}"))
    else:
        L.append(p(f"  The {project} system addresses the requirements defined in Sections 4 and 5."))

    # B.2 Data Flow Summary
    L.append(p(""))
    L.append(p("  B.2  Data Flow Summary"))
    L.append(p(f"  The {project} system processes data through the following stages:"))
    if any(k in all_text_lo for k in ["upload", "import", "input", "file"]):
        L.append(p("    1. Input: User submits data via the system interface"))
    else:
        L.append(p("    1. Input: System receives user requests"))
    L.append(p("    2. Processing: System validates and processes the data per FR specifications"))
    if any(k in all_text_lo for k in ["store", "database", "save", "record"]):
        L.append(p("    3. Storage: Processed data is persisted to the data store"))
    if any(k in all_text_lo for k in ["display", "report", "view"]):
        L.append(p("    4. Output: Results are presented to the user via the interface"))
    if any(k in all_text_lo for k in ["export", "generate", "download"]):
        L.append(p("    5. Export: Data or documents are generated for external use"))

    # B.3 Requirements Coverage Matrix
    L.append(p(""))
    L.append(p("  B.3  Requirements Coverage Matrix"))
    L.append(p(f"  Total requirements identified : {total}"))
    L.append(p(f"  Functional Requirements (FR)  : {fr_cnt}  ({metrics.get('fr_percentage', 0)}%)"))
    L.append(p(f"  Non-Functional Req.   (NFR)   : {nfr_cnt}  ({metrics.get('nfr_percentage', 0)}%)"))
    L.append(p(f"  Requirements with vague terms : {v_cnt}  ({metrics.get('vague_percentage', 0)}%)"))
    L.append(p(f"  Quality Score                 : {score} / 100  [{grade}]"))

    if vague_reqs:
        L.append(p(""))
        L.append(p("  B.4  Ambiguity Report"))
        L.append(p("  The following requirements contain vague language requiring clarification:"))
        for r in vague_reqs:
            excerpt = r['sentence'][:80] + ("..." if len(r['sentence']) > 80 else "")
            L.append(p(f"    - [{r['type']}] \"{excerpt}\""))
            L.append(p(f"      Vague terms: {', '.join(r['vague_words'])}"))

    # ── APPENDIX C: TO BE DETERMINED LIST ────────────────────────────────────
    L.append(f"\n\nAppendix C: To Be Determined List\n{SEP2}")
    if vague_reqs:
        L.append(p("The following requirements contain vague language and must be clarified:"))
        L.append(p(""))
        for idx, r in enumerate(vague_reqs, 1):
            L.append(p(f"  TBD-{idx:03d}  [{r['type']}]  {r['sentence']}"))
            L.append(p(f"           Reason: ambiguous terms — {', '.join(r['vague_words'])}"))
    else:
        L.append(p("(No TBD items — all requirements are sufficiently precise.)"))

    # ── END ──────────────────────────────────────────────────────────────────
    L += [
        f"\n\n{SEP}",
        centre("END OF SOFTWARE REQUIREMENTS SPECIFICATION"),
        SEP,
        centre("Generated from requirements specification"),
        SEP, "",
    ]

    return "\n".join(L)


def _build_docx_srs(requirements, metrics, meta, ai=None):
    """Build a FULL IEEE 830-1998 SRS as a python-docx Document, matching the TXT builder."""
    if ai is None:
        ai = {}
    try:
        from docx import Document
        from docx.shared import RGBColor, Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError(
            "python-docx is not installed. Please stop the server and run: "
            "venv\\Scripts\\pip install python-docx fpdf2"
        )

    doc = Document()

    project = meta.get("project_name", "<Project>")
    author  = meta.get("author", "<author>")
    org     = meta.get("organization", "<organization>")
    version = meta.get("version", "1.0")
    date    = meta.get("date_created") or datetime.now().strftime("%Y-%m-%d")

    frs  = [r for r in requirements if r.get("type") == "FR"]
    nfrs = [r for r in requirements if r.get("type") == "NFR"]
    total   = metrics.get("total_requirements", 0)
    fr_cnt  = metrics.get("fr_count", 0)
    nfr_cnt = metrics.get("nfr_count", 0)
    v_cnt   = metrics.get("vague_count", 0)
    score   = metrics.get("quality_score", 0)
    grade   = ("A (Excellent)" if score >= 90 else
               "B (Good)"      if score >= 75 else
               "C (Fair)"      if score >= 55 else "D (Needs Work)")

    perf_nfrs   = [r for r in nfrs if _categorise_nfr(_req_text(r)) == "performance"]
    safety_nfrs = [r for r in nfrs if _categorise_nfr(_req_text(r)) == "safety"]
    sec_nfrs    = [r for r in nfrs if _categorise_nfr(_req_text(r)) == "security"]
    qual_nfrs   = [r for r in nfrs if _categorise_nfr(_req_text(r)) == "quality"]
    other_nfrs  = [r for r in nfrs if _categorise_nfr(_req_text(r)) == "other"]
    vague_reqs  = [r for r in requirements if r.get("is_vague")]

    all_lo = " ".join(_req_text(r).lower() for r in frs + nfrs)

    # ── COVER PAGE ───────────────────────────────────────────────────
    doc.add_paragraph()  # spacer
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("Software Requirements Specification")
    run.bold = True
    run.font.size = Pt(24)

    for_p = doc.add_paragraph()
    for_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for_p.add_run("for").font.size = Pt(12)

    proj_p = doc.add_paragraph()
    proj_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = proj_p.add_run(project)
    run.bold = True
    run.font.size = Pt(18)

    ver_p = doc.add_paragraph()
    ver_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ver_p.add_run(f"Version {version} approved").font.size = Pt(11)

    auth_p = doc.add_paragraph()
    auth_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    auth_p.add_run(f"Prepared by {author}").font.size = Pt(11)

    org_p = doc.add_paragraph()
    org_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    org_p.add_run(org).font.size = Pt(11)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.add_run(date).font.size = Pt(11)

    doc.add_paragraph()
    copy_p = doc.add_paragraph()
    copy_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    copy_r = copy_p.add_run(
        "Copyright (c) 1999 by Karl E. Wiegers. "
        "Permission is granted to use, modify, and distribute this document."
    )
    copy_r.font.size = Pt(8)
    copy_r.italic = True

    doc.add_page_break()

    # ── TABLE OF CONTENTS ────────────────────────────────────────────
    doc.add_heading("Table of Contents", 1)
    toc_entries = [
        "1.  Introduction",
        "    1.1  Purpose",
        "    1.2  Document Conventions",
        "    1.3  Intended Audience and Reading Suggestions",
        "    1.4  Product Scope",
        "    1.5  References",
        "2.  Overall Description",
        "    2.1  Product Perspective",
        "    2.2  Product Functions",
        "    2.3  User Classes and Characteristics",
        "    2.4  Operating Environment",
        "    2.5  Design and Implementation Constraints",
        "    2.6  User Documentation",
        "    2.7  Assumptions and Dependencies",
        "3.  External Interface Requirements",
        "    3.1  User Interfaces",
        "    3.2  Hardware Interfaces",
        "    3.3  Software Interfaces",
        "    3.4  Communications Interfaces",
        "4.  System Features (Functional Requirements)",
        "5.  Other Nonfunctional Requirements",
        "    5.1  Performance Requirements",
        "    5.2  Safety Requirements",
        "    5.3  Security Requirements",
        "    5.4  Software Quality Attributes",
        "    5.5  Business Rules / Other NFRs",
        "6.  Other Requirements",
        "Appendix A: Glossary",
        "Appendix B: System Models",
        "Appendix C: To Be Determined List",
    ]
    for entry in toc_entries:
        doc.add_paragraph(entry)

    doc.add_page_break()

    # ── REVISION HISTORY ─────────────────────────────────────────────
    doc.add_heading("Revision History", 1)
    rev_table = doc.add_table(rows=2, cols=4)
    rev_table.style = "Light Grid Accent 1"
    for i, hdr in enumerate(["Name", "Date", "Reason For Changes", "Version"]):
        rev_table.rows[0].cells[i].text = hdr
    rev_table.rows[1].cells[0].text = "Generated"
    rev_table.rows[1].cells[1].text = date
    rev_table.rows[1].cells[2].text = "Initial SRS"
    rev_table.rows[1].cells[3].text = version

    doc.add_page_break()

    # ── §1  INTRODUCTION ─────────────────────────────────────────────
    doc.add_heading("1. Introduction", 1)

    doc.add_heading("1.1 Purpose", 2)
    doc.add_paragraph(
        f"This Software Requirements Specification (SRS) describes the requirements "
        f"for {project}. It was generated from the provided requirements document "
        f"(version {version}, {date})."
    )

    doc.add_heading("1.2 Document Conventions", 2)
    doc.add_paragraph("FR-XXX = Functional Requirement  |  NFR-XXX = Non-Functional Requirement")
    doc.add_paragraph("Statements marked [VAGUE] contain ambiguous language and must be refined.")
    doc.add_paragraph("All requirements carry equal priority unless explicitly annotated.")

    doc.add_heading("1.3 Intended Audience and Reading Suggestions", 2)
    doc.add_paragraph("Intended for developers, project managers, testers, and QA staff.")
    doc.add_paragraph("Recommendation: read Section 2 for context, then Sections 4 and 5 for specific requirements.")

    doc.add_heading("1.4 Product Scope", 2)
    scope_text = ai.get("product_scope") or _infer_scope(project, frs, nfrs)
    doc.add_paragraph(scope_text)

    doc.add_heading("1.5 References", 2)
    doc.add_paragraph("IEEE Std 830-1998, IEEE Recommended Practice for Software Requirements Specifications.")
    doc.add_paragraph("Karl E. Wiegers, Software Requirements, 2nd Edition. Microsoft Press.")
    doc.add_paragraph(f"Requirements Specification Source Document, {date}.")

    # ── §2  OVERALL DESCRIPTION ──────────────────────────────────────
    doc.add_heading("2. Overall Description", 1)

    doc.add_heading("2.1 Product Perspective", 2)
    perspective = ai.get("product_perspective") or scope_text
    doc.add_paragraph(perspective)
    doc.add_paragraph(
        f"This SRS documents {total} requirements ({fr_cnt} FR, {nfr_cnt} NFR) "
        f"covering all aspects of the {project} system."
    )

    doc.add_heading("2.2 Product Functions", 2)
    if ai.get("product_functions"):
        doc.add_paragraph(ai["product_functions"])
    else:
        doc.add_paragraph(f"Key capabilities of {project} (see Section 4 for full detail):")
        for bullet in _infer_functions(frs):
            doc.add_paragraph(bullet)

    doc.add_heading("2.3 User Classes and Characteristics", 2)
    if ai.get("user_classes"):
        doc.add_paragraph(ai["user_classes"])
    else:
        for role, desc in _infer_user_classes(frs, nfrs):
            p_uc = doc.add_paragraph()
            p_uc.add_run(f"{role}: ").bold = True
            p_uc.add_run(desc)

    doc.add_heading("2.4 Operating Environment", 2)
    if ai.get("operating_environment"):
        doc.add_paragraph(ai["operating_environment"])
    else:
        for item in _infer_environment(frs, nfrs):
            doc.add_paragraph(f"  {item}")

    doc.add_heading("2.5 Design and Implementation Constraints", 2)
    if ai.get("design_constraints"):
        doc.add_paragraph(ai["design_constraints"])
    else:
        doc.add_paragraph("The system must comply with IEEE 830-1998 specification structure.")
        doc.add_paragraph("All inter-system communication must use HTTPS with TLS 1.2 or higher.")
        doc.add_paragraph(f"{v_cnt} vague requirement(s) require clarification before implementation. See Appendix C.")

    doc.add_heading("2.6 User Documentation", 2)
    if ai.get("user_documentation"):
        doc.add_paragraph(ai["user_documentation"])
    else:
        doc.add_paragraph(f"{project} User Manual: step-by-step instructions for end users.")
        doc.add_paragraph("System Administrator Guide: installation, configuration, and maintenance.")
        doc.add_paragraph("API Reference: endpoint documentation for integration developers.")
        doc.add_paragraph("Release Notes: per-version change log and known issues.")

    doc.add_heading("2.7 Assumptions and Dependencies", 2)
    if ai.get("assumptions_dependencies"):
        doc.add_paragraph(ai["assumptions_dependencies"])
    else:
        doc.add_paragraph("The target deployment environment meets the specifications in Section 2.4.")
        if any(k in all_lo for k in ["third-party", "external", "api", "service"]):
            doc.add_paragraph("Third-party services and APIs maintain their documented SLA.")
        doc.add_paragraph("Requirements marked VAGUE will be clarified before sprint planning.")

    # ── §3  EXTERNAL INTERFACE REQUIREMENTS ──────────────────────────
    ifaces = _infer_interfaces(frs, nfrs, project)
    doc.add_heading("3. External Interface Requirements", 1)

    doc.add_heading("3.1 User Interfaces", 2)
    doc.add_paragraph(ai.get("user_interfaces") or ifaces["ui"])

    doc.add_heading("3.2 Hardware Interfaces", 2)
    doc.add_paragraph(ai.get("hardware_interfaces") or ifaces["hw"])

    doc.add_heading("3.3 Software Interfaces", 2)
    sw_text = ai.get("software_interfaces") or ifaces["sw"]
    for line in sw_text.split("\n"):
        stripped = line.strip()
        if stripped:
            doc.add_paragraph(stripped)

    doc.add_heading("3.4 Communications Interfaces", 2)
    comm_text = ai.get("communications_interfaces") or ifaces["comm"]
    for line in comm_text.split("\n"):
        stripped = line.strip()
        if stripped:
            doc.add_paragraph(stripped)

    # ── §4  SYSTEM FEATURES (Functional Requirements) ────────────────
    doc.add_heading("4. System Features (Functional Requirements)", 1)
    doc.add_paragraph(
        f"This section lists all {fr_cnt} functional requirements identified in the document. "
        f"Each feature corresponds to a requirement from the source document."
    )

    refined_frs = ai.get("refined_functional_requirements", [])
    if not refined_frs:
        refined_frs = [{"original_id": r.get("id"), "description": _req_text(r), "stimulus_response": "<Describe user actions...>"} for r in frs]

    if refined_frs:
        for i, req in enumerate(refined_frs, 1):
            doc.add_heading(f"4.{i} FR-{i:03d}", 2)

            # 4.x.1 Description and Priority
            doc.add_heading(f"4.{i}.1 Description and Priority", 3)
            doc.add_paragraph(req.get("description", ""))
            doc.add_paragraph("Priority: HIGH")

            # 4.x.2 Stimulus/Response Sequences
            doc.add_heading(f"4.{i}.2 Stimulus/Response Sequences", 3)
            doc.add_paragraph(req.get("stimulus_response", ""))

            # 4.x.3 Functional Requirements
            doc.add_heading(f"4.{i}.3 Functional Requirements", 3)
            doc.add_paragraph(f"REQ-{i:03d}: {req.get('description', '')}")
    else:
        doc.add_paragraph("(No functional requirements were identified in the source document.)")

    # ── §5  OTHER NONFUNCTIONAL REQUIREMENTS ─────────────────────────
    doc.add_heading("5. Other Nonfunctional Requirements", 1)

    refined_nfrs = ai.get("refined_non_functional_requirements", [])
    if not refined_nfrs:
        refined_nfrs = [{"original_id": r.get("id"), "description": _req_text(r), "category": "Other"} for r in nfrs]
        
    perf_nfrs   = [r for r in refined_nfrs if "performance" in r.get("category", "").lower()]
    sec_nfrs    = [r for r in refined_nfrs if "security" in r.get("category", "").lower()]
    usab_nfrs   = [r for r in refined_nfrs if "usability" in r.get("category", "").lower()]
    scal_nfrs   = [r for r in refined_nfrs if "scalability" in r.get("category", "").lower()]
    reli_nfrs   = [r for r in refined_nfrs if "reliability" in r.get("category", "").lower()]

    nfr_groups = [
        ("5.1", "Performance Requirements",  perf_nfrs, "PERF"),
        ("5.2", "Security Requirements",     sec_nfrs,  "SEC"),
        ("5.3", "Usability Requirements",    usab_nfrs, "USAB"),
        ("5.4", "Scalability Requirements",  scal_nfrs, "SCAL"),
        ("5.5", "Reliability Requirements",  reli_nfrs, "REL"),
    ]
    for sec, title_text, subset, tag in nfr_groups:
        doc.add_heading(f"{sec} {title_text}", 2)
        if subset:
            for j, r in enumerate(subset, 1):
                p_r = doc.add_paragraph()
                p_r.add_run(f"{tag}-{j:03d}: ").bold = True
                p_r.add_run(r.get("description", r.get("sentence", "")))
        else:
            p_ph = doc.add_paragraph()
            run_ph = p_ph.add_run("[PLACEHOLDER] No requirements found for this category. Please add relevant requirements or mark as Not Applicable.")
            run_ph.font.color.rgb = RGBColor(0xFF, 0x8C, 0x00)
            run_ph.italic = True

    # ── §6  OTHER REQUIREMENTS ───────────────────────────────────────
    doc.add_heading("6. Other Requirements", 1)
    if ai.get("other_requirements"):
        doc.add_paragraph(ai["other_requirements"])
    else:
        if any(k in all_lo for k in ["language", "locale", "internation", "multilingual"]):
            doc.add_paragraph(
                "Internationalisation: The system shall support multiple locales/languages "
                "as specified during project initiation."
            )
        if any(k in all_lo for k in ["gdpr", "legal", "compliance", "regulation"]):
            doc.add_paragraph(
                "Legal/Regulatory: The system must comply with applicable data protection "
                "regulations (e.g., GDPR, HIPAA) as required by jurisdiction."
            )
        if any(k in all_lo for k in ["database", "backup", "archive", "retention"]):
            doc.add_paragraph(
                "Data Retention: Data shall be retained for a minimum period as defined "
                "by applicable policy; purged securely thereafter."
            )
        else:
            doc.add_paragraph("Database: Data persistence and storage requirements are detailed in Section 3.3.")
            doc.add_paragraph(
                "Legal: The system shall comply with applicable laws and regulations in the "
                "target deployment region."
            )

    # ── APPENDIX A: GLOSSARY ─────────────────────────────────────────
    doc.add_heading("Appendix A: Glossary", 1)
    glossary = [
        ("FR",    "Functional Requirement: a capability the system must perform."),
        ("NFR",   "Non-Functional Requirement: a quality/constraint the system must meet."),
        ("VAGUE", "Ambiguous language; statement needs a quantifiable, measurable form."),
        ("TBD",   "To Be Determined; information not yet available."),
        ("SRS",   "Software Requirements Specification."),
        ("IEEE",  "Institute of Electrical and Electronics Engineers."),
    ]
    if any(k in all_lo for k in ["api", "rest"]):
        glossary.append(("API", "Application Programming Interface: a set of defined interaction rules."))
    if any(k in all_lo for k in ["authenticate", "auth", "login"]):
        glossary.append(("Auth", "Authentication/Authorization: verifying identity and access rights."))

    for term, defn in glossary:
        p_gl = doc.add_paragraph()
        p_gl.add_run(f"{term}").bold = True
        p_gl.add_run(f" - {defn}")

    # ── APPENDIX B: ANALYSIS MODELS ──────────────────────────────────
    doc.add_heading("Appendix B: System Models", 1)
    doc.add_paragraph(f"System Models for {project}")

    # B.1 System Overview
    doc.add_heading("B.1 System Overview", 2)
    all_text_b = " ".join(_req_text(r).lower() for r in frs + nfrs)
    themes_bd = []
    theme_map_bd = [
        (["upload", "import", "file", "document"], "Document Ingestion and Processing"),
        (["login", "auth", "register", "account"], "User Authentication and Account Management"),
        (["search", "filter", "query", "browse"], "Search and Data Retrieval"),
        (["report", "analytics", "chart"], "Reporting and Analytics"),
        (["export", "generate", "download"], "Output and Data Export"),
        (["notify", "alert", "email", "message"], "Notifications and Messaging"),
        (["payment", "invoice", "billing"], "Payment Processing"),
        (["admin", "manage", "configure", "role"], "Administration and Configuration"),
    ]
    for kws_bd, lbl_bd in theme_map_bd:
        if any(k in all_text_b for k in kws_bd):
            themes_bd.append(lbl_bd)
    if themes_bd:
        doc.add_paragraph(f"The {project} system provides the following core capabilities:")
        for tbd in themes_bd:
            doc.add_paragraph(f"  - {tbd}")
    else:
        doc.add_paragraph(f"The {project} system addresses the requirements defined in Sections 4 and 5.")

    # B.2 Data Flow Summary
    doc.add_heading("B.2 Data Flow Summary", 2)
    doc.add_paragraph(f"The {project} system processes data through the following stages:")
    flow_steps = []
    if any(k in all_text_b for k in ["upload", "import", "input", "file"]):
        flow_steps.append("1. Input: User submits data via the system interface")
    else:
        flow_steps.append("1. Input: System receives user requests")
    flow_steps.append("2. Processing: System validates and processes the data per FR specifications")
    if any(k in all_text_b for k in ["store", "database", "save", "record"]):
        flow_steps.append("3. Storage: Processed data is persisted to the data store")
    if any(k in all_text_b for k in ["display", "report", "view"]):
        flow_steps.append("4. Output: Results are presented to the user via the interface")
    if any(k in all_text_b for k in ["export", "generate", "download"]):
        flow_steps.append("5. Export: Data or documents are generated for external use")
    for step in flow_steps:
        doc.add_paragraph(step)

    # B.3 Requirements Coverage Matrix
    doc.add_heading("B.3 Requirements Coverage Matrix", 2)
    cov_table = doc.add_table(rows=5, cols=2)
    cov_table.style = "Light Grid Accent 1"
    cov_rows = [
        ("Total Requirements", str(total)),
        ("Functional (FR)", f"{fr_cnt}  ({metrics.get('fr_percentage', 0)}%)"),
        ("Non-Functional (NFR)", f"{nfr_cnt}  ({metrics.get('nfr_percentage', 0)}%)"),
        ("Requirements with Vague Terms", f"{v_cnt}  ({metrics.get('vague_percentage', 0)}%)"),
        ("Quality Score", f"{score} / 100  [{grade}]"),
    ]
    for ci, (lb, vl) in enumerate(cov_rows):
        cov_table.rows[ci].cells[0].text = lb
        cov_table.rows[ci].cells[1].text = vl

    if vague_reqs:
        doc.add_paragraph()
        doc.add_heading("B.4 Ambiguity Report", 2)
        doc.add_paragraph("The following requirements contain vague language requiring clarification:")
        for r in vague_reqs:
            req_text = _req_text(r)
            excerpt = req_text[:80] + ("..." if len(req_text) > 80 else "")
            p_amb = doc.add_paragraph()
            p_amb.add_run(f"[{r['type']}] ").bold = True
            p_amb.add_run(f'"{excerpt}"')
            doc.add_paragraph(f"    Vague terms: {', '.join(r.get('vague_words', []))}")

    # ── APPENDIX C: TO BE DETERMINED LIST ────────────────────────────
    doc.add_heading("Appendix C: To Be Determined List", 1)
    if vague_reqs:
        doc.add_paragraph("The following requirements contain vague language and must be clarified:")
        for idx, r in enumerate(vague_reqs, 1):
            p_tbd = doc.add_paragraph()
            p_tbd.add_run(f"TBD-{idx:03d}  [{r['type']}]  ").bold = True
            p_tbd.add_run(_req_text(r))
            doc.add_paragraph(f"    Reason: ambiguous terms - {', '.join(r.get('vague_words', []))}")
    else:
        doc.add_paragraph("(No TBD items - all requirements are sufficiently precise.)")

    # ── END ──────────────────────────────────────────────────────────
    doc.add_paragraph()
    end_p = doc.add_paragraph()
    end_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    end_r = end_p.add_run("END OF SOFTWARE REQUIREMENTS SPECIFICATION")
    end_r.bold = True

    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_r = footer_p.add_run("Generated from requirements specification")
    footer_r.italic = True
    footer_r.font.size = Pt(9)

    return doc


def _build_pdf_srs(requirements, metrics, meta, ai=None) -> bytes:
    """Build a FULL IEEE 830-1998 SRS as a PDF matching the TXT/DOCX builders."""
    if ai is None:
        ai = {}
    try:
        from fpdf import FPDF
    except ImportError:
        raise ImportError(
            "fpdf2 is not installed. Please stop the server and run: "
            "venv\\Scripts\\pip install python-docx fpdf2"
        )

    project = meta.get("project_name", "<Project>")
    author  = meta.get("author", "<author>")
    org     = meta.get("organization", "<organization>")
    version = meta.get("version", "1.0")
    date    = meta.get("date_created") or datetime.now().strftime("%Y-%m-%d")

    total   = metrics.get("total_requirements", 0)
    fr_cnt  = metrics.get("fr_count", 0)
    nfr_cnt = metrics.get("nfr_count", 0)
    v_cnt   = metrics.get("vague_count", 0)
    score   = metrics.get("quality_score", 0)
    grade   = ("A (Excellent)" if score >= 90 else
               "B (Good)"      if score >= 75 else
               "C (Fair)"      if score >= 55 else "D (Needs Work)")

    frs  = [r for r in requirements if r.get("type") == "FR"]
    nfrs = [r for r in requirements if r.get("type") == "NFR"]
    vague_reqs = [r for r in requirements if r.get("is_vague")]
    all_lo = " ".join(_req_text(r).lower() for r in frs + nfrs)

    perf_nfrs   = [r for r in nfrs if _categorise_nfr(_req_text(r)) == "performance"]
    safety_nfrs = [r for r in nfrs if _categorise_nfr(_req_text(r)) == "safety"]
    sec_nfrs    = [r for r in nfrs if _categorise_nfr(_req_text(r)) == "security"]
    qual_nfrs   = [r for r in nfrs if _categorise_nfr(_req_text(r)) == "quality"]
    other_nfrs  = [r for r in nfrs if _categorise_nfr(_req_text(r)) == "other"]

    # ── Sanitize text for built-in Helvetica (latin-1 only) ──
    def _s(text):
        t = str(text)
        for old, new in [
            ('\u2022', '*'), ('\u2013', '-'), ('\u2014', '--'),
            ('\u2018', "'"), ('\u2019', "'"), ('\u201c', '"'), ('\u201d', '"'),
            ('\u2265', '>='), ('\u2264', '<='), ('\u00a7', 'S.'),
            ('\u2026', '...'), ('\u00b7', '*'), ('\u2192', '->'),
            ('\u26a0', '[!]'), ('\u2713', '[v]'), ('\u2717', '[x]'),
            ('\u2611', '[x]'), ('\u2610', '[ ]'),
        ]:
            t = t.replace(old, new)
        t = t.encode('latin-1', errors='replace').decode('latin-1')
        return t

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)

    def section_heading(num, title, level=1):
        if level == 1:
            pdf.set_font("Helvetica", "B", 13)
            pdf.set_fill_color(240, 240, 255)
            pdf.cell(0, 8, _s(f"{num}  {title}"), new_x="LMARGIN", new_y="NEXT", fill=True)
            pdf.ln(2)
        elif level == 2:
            pdf.set_font("Helvetica", "B", 11)
            pdf.cell(0, 7, _s(f"{num}  {title}"), new_x="LMARGIN", new_y="NEXT")
        else:
            pdf.set_font("Helvetica", "B", 10)
            pdf.cell(0, 6, _s(f"{num}  {title}"), new_x="LMARGIN", new_y="NEXT")

    def body_text(text, indent=4):
        pdf.set_font("Helvetica", "", 9)
        pdf.set_x(10 + indent)
        pdf.multi_cell(0, 5, _s(text))
        pdf.ln(1)

    def vague_text(text):
        pdf.set_font("Helvetica", "I", 8)
        pdf.set_text_color(200, 130, 0)
        pdf.set_x(14)
        pdf.multi_cell(0, 4, _s(f"  [!] VAGUE: {text}"))
        pdf.set_text_color(0, 0, 0)

    # ── COVER PAGE ───────────────────────────────────────────────────
    pdf.add_page()
    pdf.ln(40)
    pdf.set_font("Helvetica", "B", 24)
    pdf.cell(0, 12, "Software Requirements Specification", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)
    pdf.set_font("Helvetica", "", 12)
    pdf.cell(0, 8, "for", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(2)
    pdf.set_font("Helvetica", "B", 18)
    pdf.cell(0, 10, _s(project), align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(6)
    pdf.set_font("Helvetica", "", 11)
    pdf.cell(0, 6, _s(f"Version {version} approved"), align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(2)
    pdf.cell(0, 6, _s(f"Prepared by {author}"), align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.cell(0, 6, _s(org), align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.cell(0, 6, _s(date), align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(12)
    pdf.set_draw_color(99, 102, 241)
    pdf.set_line_width(0.5)
    pdf.line(30, pdf.get_y(), 180, pdf.get_y())
    pdf.ln(6)
    pdf.set_font("Helvetica", "I", 8)
    pdf.cell(0, 5, "Copyright (c) 1999 by Karl E. Wiegers. Permission is granted to use, modify, and distribute this document.",
             align="C", new_x="LMARGIN", new_y="NEXT")

    # ── TABLE OF CONTENTS ────────────────────────────────────────────
    pdf.add_page()
    section_heading("", "Table of Contents")
    toc_entries = [
        "1.  Introduction",
        "    1.1  Purpose",
        "    1.2  Document Conventions",
        "    1.3  Intended Audience and Reading Suggestions",
        "    1.4  Product Scope",
        "    1.5  References",
        "2.  Overall Description",
        "    2.1  Product Perspective",
        "    2.2  Product Functions",
        "    2.3  User Classes and Characteristics",
        "    2.4  Operating Environment",
        "    2.5  Design and Implementation Constraints",
        "    2.6  User Documentation",
        "    2.7  Assumptions and Dependencies",
        "3.  External Interface Requirements",
        "4.  System Features (Functional Requirements)",
        "5.  Other Nonfunctional Requirements",
        "6.  Other Requirements",
        "Appendix A: Glossary",
        "Appendix B: System Models",
        "Appendix C: To Be Determined List",
    ]
    for entry in toc_entries:
        body_text(entry, indent=2)

    # ── REVISION HISTORY ─────────────────────────────────────────────
    pdf.add_page()
    section_heading("", "Revision History")
    pdf.set_font("Helvetica", "B", 9)
    col_w = [35, 25, 85, 20]
    for w, h in zip(col_w, ["Name", "Date", "Reason For Changes", "Version"]):
        pdf.cell(w, 6, h, border=1)
    pdf.ln()
    pdf.set_font("Helvetica", "", 9)
    for w, val in zip(col_w, ["Generated", date, "Initial SRS", version]):
        pdf.cell(w, 6, _s(val), border=1)
    pdf.ln(8)

    # ── S.1  INTRODUCTION ────────────────────────────────────────────
    pdf.add_page()
    section_heading("1.", "Introduction")

    section_heading("1.1", "Purpose", 2)
    body_text(f"This Software Requirements Specification (SRS) describes the requirements "
              f"for {project}. It was generated from the provided requirements document "
              f"(version {version}, {date}).")

    section_heading("1.2", "Document Conventions", 2)
    body_text("FR-XXX = Functional Requirement  |  NFR-XXX = Non-Functional Requirement")
    body_text("Statements marked [VAGUE] contain ambiguous language and must be refined.")
    body_text("All requirements carry equal priority unless explicitly annotated.")

    section_heading("1.3", "Intended Audience and Reading Suggestions", 2)
    body_text("Intended for developers, project managers, testers, and QA staff.")
    body_text("Recommendation: read Section 2 for context, then Sections 4 and 5 for specific requirements.")

    section_heading("1.4", "Product Scope", 2)
    scope_text = ai.get("product_scope") or _infer_scope(project, frs, nfrs)
    body_text(scope_text)

    section_heading("1.5", "References", 2)
    body_text("IEEE Std 830-1998, IEEE Recommended Practice for Software Requirements Specifications.")
    body_text("Karl E. Wiegers, Software Requirements, 2nd Edition. Microsoft Press.")
    body_text(f"Requirements Specification Source Document, {date}.")

    # ── S.2  OVERALL DESCRIPTION ─────────────────────────────────────
    section_heading("2.", "Overall Description")

    section_heading("2.1", "Product Perspective", 2)
    perspective = ai.get("product_perspective") or scope_text
    body_text(perspective)
    body_text(f"This SRS documents {total} requirements ({fr_cnt} FR, {nfr_cnt} NFR) "
              f"covering all aspects of the {project} system.")

    section_heading("2.2", "Product Functions", 2)
    if ai.get("product_functions"):
        body_text(ai["product_functions"])
    else:
        body_text(f"Key capabilities of {project} (see Section 4 for full detail):")
        for bullet in _infer_functions(frs):
            body_text(f"  {bullet}")

    section_heading("2.3", "User Classes and Characteristics", 2)
    if ai.get("user_classes"):
        body_text(ai["user_classes"])
    else:
        for role, desc in _infer_user_classes(frs, nfrs):
            body_text(f"{role}: {desc}")

    section_heading("2.4", "Operating Environment", 2)
    if ai.get("operating_environment"):
        body_text(ai["operating_environment"])
    else:
        for item in _infer_environment(frs, nfrs):
            body_text(f"  * {item}")

    section_heading("2.5", "Design and Implementation Constraints", 2)
    if ai.get("design_constraints"):
        body_text(ai["design_constraints"])
    else:
        body_text("The system must comply with IEEE 830-1998 specification structure.")
        body_text("All inter-system communication must use HTTPS with TLS 1.2 or higher.")
        body_text(f"{v_cnt} vague requirement(s) require clarification before implementation. See Appendix C.")

    section_heading("2.6", "User Documentation", 2)
    if ai.get("user_documentation"):
        body_text(ai["user_documentation"])
    else:
        body_text(f"{project} User Manual: step-by-step instructions for end users.")
        body_text("System Administrator Guide: installation, configuration, and maintenance.")
        body_text("API Reference: endpoint documentation for integration developers.")
        body_text("Release Notes: per-version change log and known issues.")

    section_heading("2.7", "Assumptions and Dependencies", 2)
    if ai.get("assumptions_dependencies"):
        body_text(ai["assumptions_dependencies"])
    else:
        body_text("The target deployment environment meets the specifications in Section 2.4.")
        if any(k in all_lo for k in ["third-party", "external", "api", "service"]):
            body_text("Third-party services and APIs maintain their documented SLA.")
        body_text("Requirements marked VAGUE will be clarified before sprint planning.")

    # ── S.3  EXTERNAL INTERFACE REQUIREMENTS ─────────────────────────
    ifaces = _infer_interfaces(frs, nfrs, project)
    section_heading("3.", "External Interface Requirements")

    section_heading("3.1", "User Interfaces", 2)
    body_text(ai.get("user_interfaces") or ifaces["ui"])

    section_heading("3.2", "Hardware Interfaces", 2)
    body_text(ai.get("hardware_interfaces") or ifaces["hw"])

    section_heading("3.3", "Software Interfaces", 2)
    sw_text = ai.get("software_interfaces") or ifaces["sw"]
    for line in sw_text.split("\n"):
        stripped = line.strip()
        if stripped:
            body_text(stripped)

    section_heading("3.4", "Communications Interfaces", 2)
    comm_text = ai.get("communications_interfaces") or ifaces["comm"]
    for line in comm_text.split("\n"):
        stripped = line.strip()
        if stripped:
            body_text(stripped)

    # ── S.4  SYSTEM FEATURES (Functional Requirements) ───────────────
    section_heading("4.", "System Features (Functional Requirements)")
    body_text(f"This section lists all {fr_cnt} functional requirements identified in the document. "
              f"Each feature corresponds to a requirement from the source document.")

    refined_frs = ai.get("refined_functional_requirements", [])
    if not refined_frs:
        refined_frs = [{"original_id": r.get("id"), "description": _req_text(r), "stimulus_response": "<Describe user actions...>"} for r in frs]

    if refined_frs:
        for i, req in enumerate(refined_frs, 1):
            section_heading(f"4.{i}", f"FR-{i:03d}", 2)

            section_heading(f"4.{i}.1", "Description and Priority", 3)
            body_text(req.get("description", ""))
            body_text("Priority: HIGH")

            section_heading(f"4.{i}.2", "Stimulus/Response Sequences", 3)
            body_text(req.get("stimulus_response", ""))

            section_heading(f"4.{i}.3", "Functional Requirements", 3)
            body_text(f"REQ-{i:03d}: {req.get('description', '')}")
    else:
        body_text("(No functional requirements were identified in the source document.)")

    # ── S.5  OTHER NONFUNCTIONAL REQUIREMENTS ────────────────────────
    section_heading("5.", "Other Nonfunctional Requirements")

    refined_nfrs = ai.get("refined_non_functional_requirements", [])
    if not refined_nfrs:
        refined_nfrs = [{"original_id": r.get("id"), "description": _req_text(r), "category": "Other"} for r in nfrs]
        
    perf_nfrs   = [r for r in refined_nfrs if "performance" in r.get("category", "").lower()]
    sec_nfrs    = [r for r in refined_nfrs if "security" in r.get("category", "").lower()]
    usab_nfrs   = [r for r in refined_nfrs if "usability" in r.get("category", "").lower()]
    scal_nfrs   = [r for r in refined_nfrs if "scalability" in r.get("category", "").lower()]
    reli_nfrs   = [r for r in refined_nfrs if "reliability" in r.get("category", "").lower()]

    nfr_groups = [
        ("5.1", "Performance Requirements",  perf_nfrs, "PERF"),
        ("5.2", "Security Requirements",     sec_nfrs,  "SEC"),
        ("5.3", "Usability Requirements",    usab_nfrs, "USAB"),
        ("5.4", "Scalability Requirements",  scal_nfrs, "SCAL"),
        ("5.5", "Reliability Requirements",  reli_nfrs, "REL"),
    ]
    for sec, title_text, subset, tag in nfr_groups:
        section_heading(sec, title_text, 2)
        if subset:
            for j, r in enumerate(subset, 1):
                body_text(f"{tag}-{j:03d}: {r.get('description', r.get('sentence', ''))}")
        else:
            pdf.set_text_color(255, 140, 0)
            pdf.set_font("Helvetica", "I", 9)
            pdf.multi_cell(0, 5, _s("[PLACEHOLDER] No requirements found for this category. Please add relevant requirements or mark as Not Applicable."))
            pdf.set_text_color(60, 60, 60)
            pdf.set_font("Helvetica", "", 10)
            pdf.ln(2)

    # ── S.6  OTHER REQUIREMENTS ──────────────────────────────────────
    section_heading("6.", "Other Requirements")
    if ai.get("other_requirements"):
        body_text(ai["other_requirements"])
    else:
        if any(k in all_lo for k in ["language", "locale", "internation", "multilingual"]):
            body_text("Internationalisation: The system shall support multiple locales/languages "
                      "as specified during project initiation.")
        if any(k in all_lo for k in ["gdpr", "legal", "compliance", "regulation"]):
            body_text("Legal/Regulatory: The system must comply with applicable data protection "
                      "regulations (e.g., GDPR, HIPAA) as required by jurisdiction.")
        if any(k in all_lo for k in ["database", "backup", "archive", "retention"]):
            body_text("Data Retention: Data shall be retained for a minimum period as defined "
                      "by applicable policy; purged securely thereafter.")
        else:
            body_text("Database: Data persistence and storage requirements are detailed in Section 3.3.")
            body_text("Legal: The system shall comply with applicable laws and regulations in the "
                      "target deployment region.")

    # ── APPENDIX A: GLOSSARY ─────────────────────────────────────────
    section_heading("Appendix A:", "Glossary")
    glossary = [
        ("FR",    "Functional Requirement: a capability the system must perform."),
        ("NFR",   "Non-Functional Requirement: a quality/constraint the system must meet."),
        ("VAGUE", "Ambiguous language; statement needs a quantifiable, measurable form."),
        ("TBD",   "To Be Determined; information not yet available."),
        ("SRS",   "Software Requirements Specification."),
        ("IEEE",  "Institute of Electrical and Electronics Engineers."),
    ]
    if any(k in all_lo for k in ["api", "rest"]):
        glossary.append(("API", "Application Programming Interface."))
    if any(k in all_lo for k in ["authenticate", "auth", "login"]):
        glossary.append(("Auth", "Authentication/Authorization."))
    for term, defn in glossary:
        body_text(f"{term} -- {defn}")

    # ── APPENDIX B: ANALYSIS MODELS ──────────────────────────────────
    section_heading("Appendix B:", "System Models")
    body_text(f"System Models for {project}")

    # B.1 System Overview
    section_heading("B.1", "System Overview", 2)
    all_text_bp = " ".join(r.get("sentence", r.get("description", "")).lower() for r in frs + nfrs)
    themes_bp = []
    theme_map_bp = [
        (["upload", "import", "file", "document"], "Document Ingestion and Processing"),
        (["login", "auth", "register", "account"], "User Authentication and Account Management"),
        (["search", "filter", "query", "browse"], "Search and Data Retrieval"),
        (["report", "analytics", "chart"], "Reporting and Analytics"),
        (["export", "generate", "download"], "Output and Data Export"),
        (["notify", "alert", "email", "message"], "Notifications and Messaging"),
        (["payment", "invoice", "billing"], "Payment Processing"),
        (["admin", "manage", "configure", "role"], "Administration and Configuration"),
    ]
    for kws_bp, lbl_bp in theme_map_bp:
        if any(k in all_text_bp for k in kws_bp):
            themes_bp.append(lbl_bp)
    if themes_bp:
        body_text(f"The {project} system provides the following core capabilities:")
        for tbp in themes_bp:
            body_text(f"  - {tbp}")
    else:
        body_text(f"The {project} system addresses the requirements defined in Sections 4 and 5.")

    # B.2 Data Flow
    section_heading("B.2", "Data Flow Summary", 2)
    body_text(f"The {project} system processes data through the following stages:")
    if any(k in all_text_bp for k in ["upload", "import", "input", "file"]):
        body_text("1. Input: User submits data via the system interface")
    else:
        body_text("1. Input: System receives user requests")
    body_text("2. Processing: System validates and processes the data per FR specifications")
    if any(k in all_text_bp for k in ["store", "database", "save", "record"]):
        body_text("3. Storage: Processed data is persisted to the data store")
    if any(k in all_text_bp for k in ["display", "report", "view"]):
        body_text("4. Output: Results are presented to the user via the interface")
    if any(k in all_text_bp for k in ["export", "generate", "download"]):
        body_text("5. Export: Data or documents are generated for external use")

    # B.3 Requirements Coverage
    section_heading("B.3", "Requirements Coverage Matrix", 2)
    pdf.set_font("Helvetica", "B", 9)
    col_w2 = [90, 60]
    for w, h in zip(col_w2, ["Metric", "Value"]):
        pdf.cell(w, 6, h, border=1)
    pdf.ln()
    pdf.set_font("Helvetica", "", 9)
    cov_rows_p = [
        ("Total Requirements", str(total)),
        ("Functional (FR)", f"{fr_cnt}  ({metrics.get('fr_percentage', 0)}%)"),
        ("Non-Functional (NFR)", f"{nfr_cnt}  ({metrics.get('nfr_percentage', 0)}%)"),
        ("Requirements with Vague Terms", f"{v_cnt}  ({metrics.get('vague_percentage', 0)}%)"),
        ("Quality Score", f"{score} / 100  [{grade}]"),
    ]
    for label, val in cov_rows_p:
        pdf.cell(col_w2[0], 6, _s(label), border=1)
        pdf.cell(col_w2[1], 6, _s(val), border=1)
        pdf.ln()
    pdf.ln(4)
    pdf.set_font("Helvetica", "", 10)

    if vague_reqs:
        section_heading("B.4", "Ambiguity Report", 2)
        body_text("The following requirements contain vague language requiring clarification:")
        for r in vague_reqs:
            req_text = _req_text(r)
            excerpt = req_text[:80] + ("..." if len(req_text) > 80 else "")
            body_text(f'[{r["type"]}] "{excerpt}"')
            body_text(f"  Vague terms: {', '.join(r.get('vague_words', []))}", indent=8)

    # ── APPENDIX C: TO BE DETERMINED LIST ────────────────────────────
    section_heading("Appendix C:", "To Be Determined List")
    if vague_reqs:
        body_text("The following requirements contain vague language and must be clarified:")
        for idx, r in enumerate(vague_reqs, 1):
            body_text(f"TBD-{idx:03d}  [{r['type']}]  {_req_text(r)}")
            body_text(f"  Reason: ambiguous terms - {', '.join(r.get('vague_words', []))}", indent=8)
    else:
        body_text("(No TBD items - all requirements are sufficiently precise.)")

    # ── END ──────────────────────────────────────────────────────────
    pdf.ln(6)
    pdf.set_draw_color(99, 102, 241)
    pdf.line(30, pdf.get_y(), 180, pdf.get_y())
    pdf.ln(4)
    pdf.set_font("Helvetica", "B", 10)
    pdf.cell(0, 6, "END OF SOFTWARE REQUIREMENTS SPECIFICATION", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(2)
    pdf.set_font("Helvetica", "I", 8)
    pdf.set_text_color(120, 120, 120)
    pdf.cell(0, 5, _s("Generated from requirements specification"),
             align="C", new_x="LMARGIN", new_y="NEXT")

    return bytes(pdf.output())



# ─────────────────────────────────────────────────────────────────────────────
# System Boundary Enforcement
# ─────────────────────────────────────────────────────────────────────────────

# Phrases that belong to ARAQAT (the tool), NOT the target system
_TOOL_PHRASES = [
    "nlp", "natural language processing", "requirement extraction",
    "extract requirement", "automated srs generation", "generate srs",
    "quality scoring", "quality assessment", "ambiguity detection",
    "analysis dashboard", "analyzes documents", "analyze documents",
    "araqat", "automated requirements analysis",
]


def _is_tool_specific(text: str) -> bool:
    """Return True if the text describes ARAQAT capabilities, not the target system."""
    lo = (text or "").lower()
    return any(phrase in lo for phrase in _TOOL_PHRASES)


def _filter_target_system_requirements(refined: list) -> list:
    """
    System Boundary Enforcement — strip any refined requirements that describe
    ARAQAT's own capabilities (NLP, SRS generation, quality scoring, etc.)
    instead of the target system's features.

    These should never appear in the SRS — they belong to the tool, not the
    system being specified.
    """
    filtered = []
    for req in refined:
        text = req.get("text", req.get("sentence", req.get("description", "")))
        if _is_tool_specific(text):
            print(f"[boundary] Removed tool-specific requirement: {text[:80]}")
            continue
        filtered.append(req)
    return filtered


def _sanitize_ai_content(ai_content: dict) -> dict:
    """
    Scrub ARAQAT-specific phrases from LLM-generated prose sections.

    The LLM sometimes echoes tool capabilities (e.g. 'NLP-powered analysis',
    'automated SRS generation') into the SRS prose because they appear in the
    prompt context. This function detects and neutralises such phrases.
    """
    if not ai_content:
        return ai_content

    # Replacement map: tool phrase → neutral alternative
    REPLACEMENTS = [
        ("nlp-powered", "automated"),
        ("nlp based", "automated"),
        ("nlp analysis", "automated analysis"),
        ("natural language processing", "automated text analysis"),
        ("requirement extraction", "requirement identification"),
        ("automated srs generation", "document generation"),
        ("generate srs", "generate documentation"),
        ("quality scoring", "quality evaluation"),
        ("quality assessment", "quality evaluation"),
        ("ambiguity detection", "requirement review"),
        ("analysis dashboard", "reporting dashboard"),
        ("analyzes documents", "processes documents"),
        ("araqat", "the system"),
        ("automated requirements analysis and quality assessment tool", "the system"),
    ]

    sanitized = {}
    for key, value in ai_content.items():
        if isinstance(value, str):
            for phrase, replacement in REPLACEMENTS:
                import re as _re
                value = _re.sub(_re.escape(phrase), replacement, value, flags=_re.IGNORECASE)
            sanitized[key] = value
        else:
            sanitized[key] = value
    return sanitized


@app.route("/api/export", methods=["POST"])
def export_srs():
    """
    Accept analyzed JSON + project metadata from the frontend.
    Pipeline: Extract → Clean → Refine (LLM 1) → Generate Prose (LLM 2)
    Returns an IEEE 830-1998 SRS in txt, docx, or pdf format.
    Query param: ?format=txt|docx|pdf (default: txt)
    """
    fmt = request.args.get("format", "txt").lower()
    data = request.get_json()
    if not data:
        return jsonify({"error": "No data provided."}), 400

    raw_requirements = data.get("requirements", [])
    metrics          = data.get("metrics", {})
    meta             = {
        "project_name": data.get("project_name", "<Project>"),
        "author":       data.get("author",       "<author>"),
        "organization": data.get("organization", "<organization>"),
        "version":      data.get("version",      "1.0"),
        "date_created": data.get("date_created", datetime.now().strftime("%Y-%m-%d")),
    }
    safe_name = (meta["project_name"] or "Project").replace(" ", "_")

    # ── STAGE 1: CLEAN raw requirements (returns dict with 'requirements' + 'routed') ──
    clean_result = clean_requirements(raw_requirements)
    cleaned = clean_result["requirements"]
    routed  = clean_result["routed"]

    # ── STAGE 2: REFINE via LLM Call 1 (or fallback) ──
    refined = refine_requirements(cleaned, meta)
    refined = _filter_target_system_requirements(refined)

    # ── STAGE 3: GENERATE section prose via LLM Call 2 (or fallback) ──
    ai_content = generate_srs_content(refined, meta)
    if ai_content is None:
        ai_content = generate_fallback_content(refined, meta)
    ai_content = _sanitize_ai_content(ai_content)

    # Integrate routed non-requirement content into SRS prose
    if routed.get("scope_hints"):
        existing_scope = ai_content.get("product_scope", "")
        extra = " ".join(routed["scope_hints"])
        ai_content["product_scope"] = f"{existing_scope}\n\nProject Context: {extra}"
    if routed.get("future_scope"):
        existing_assumptions = ai_content.get("assumptions_dependencies", "")
        future_items = "\n".join(f"  - Future scope: {s}" for s in routed["future_scope"])
        ai_content["assumptions_dependencies"] = f"{existing_assumptions}\n{future_items}"
    if routed.get("environment_hints"):
        existing_env = ai_content.get("operating_environment", "")
        env_items = "\n".join(f"  - {s}" for s in routed["environment_hints"])
        ai_content["operating_environment"] = f"{existing_env}\n{env_items}"
    if routed.get("incomplete"):
        tbd_items = "\n".join(
            f"  - [{item['original_id']}] {item['sentence']} ({item['flag']})"
            for item in routed["incomplete"]
        )
        existing_other = ai_content.get("other_requirements", "")
        ai_content["other_requirements"] = f"{existing_other}\n\nIncomplete Requirements:\n{tbd_items}"
    ai_content = _sanitize_ai_content(ai_content)

    # Attach the refined requirements into ai_content so renderers can use them
    frs  = [r for r in refined if r.get("type") == "FR"]
    nfrs = [r for r in refined if r.get("type") == "NFR"]
    ai_content["refined_functional_requirements"] = [
        {
            "original_id": r.get("id", f"FR-{i+1:03d}"),
            "description": r.get("text", ""),
            "stimulus_response": r.get("stimulus_response", ""),
            "original_text": r.get("original", ""),
        }
        for i, r in enumerate(frs)
    ]
    ai_content["refined_non_functional_requirements"] = [
        {
            "original_id": r.get("id", f"NFR-{i+1:03d}"),
            "description": r.get("text", ""),
            "category": r.get("category", "Performance"),
            "original_text": r.get("original", ""),
        }
        for i, r in enumerate(nfrs)
    ]

    # Update metrics counts based on refined data
    metrics["fr_count"] = len(frs)
    metrics["nfr_count"] = len(nfrs)
    metrics["total_requirements"] = len(refined)

    try:
        if fmt == "docx":
            doc = _build_docx_srs(refined, metrics, meta, ai=ai_content)
            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            return send_file(
                buf,
                mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                as_attachment=True,
                download_name=f"SRS_{safe_name}_IEEE830.docx",
            )
        elif fmt == "pdf":
            pdf_bytes = _build_pdf_srs(refined, metrics, meta, ai=ai_content)
            buf = io.BytesIO(pdf_bytes)
            buf.seek(0)
            return send_file(
                buf,
                mimetype="application/pdf",
                as_attachment=True,
                download_name=f"SRS_{safe_name}_IEEE830.pdf",
            )
        else:
            # Default: TXT — return JSON with srs key for preview
            srs_text = _build_ieee_srs(refined, metrics, meta, ai=ai_content)
            return jsonify({"srs": srs_text}), 200

    except Exception as e:
        return jsonify({"error": f"SRS generation failed: {str(e)}"}), 500


if __name__ == "__main__":
    app.run(debug=True, port=5000)

